import os
import base64
import io
import json
import re
import traceback
from flask import Blueprint, request, jsonify, send_file
from openai import OpenAI, APIConnectionError
from PIL import Image as PILImage
import requests
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from xml.sax.saxutils import escape
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.graphics import renderPDF
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from docx import Document
from docx.shared import Inches, Pt
import matplotlib
from matplotlib import rcParams
matplotlib.use('Agg')
rcParams['mathtext.fontset'] = 'stix'
rcParams['font.family'] = 'STIXGeneral'
rcParams['mathtext.default'] = 'regular'
from matplotlib.mathtext import math_to_image
from matplotlib.font_manager import FontProperties

try:
    from svglib.svglib import svg2rlg
except ImportError:
    svg2rlg = None

math_bp = Blueprint('math', __name__)

DEFAULT_FONT_NAME = 'HeiseiKakuGo-W5'
try:
    pdfmetrics.registerFont(UnicodeCIDFont(DEFAULT_FONT_NAME))
except Exception:
    DEFAULT_FONT_NAME = 'Helvetica'
# OpenAI クライアントの初期化
client = OpenAI()

def load_prompt_template():
    template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'prompt_template.txt')
    with open(template_path, 'r', encoding='utf-8') as f:
        return f.read()


MATH_PATTERN = re.compile(
    r'(?<!\\)(?:\$\$.*?(?<!\\)\$\$|\$(?!\s).*?(?<!\\)\$|\\\(.*?(?<!\\)\\\)|\\\[.*?(?<!\\)\\\])',
    re.DOTALL,
)




def split_text_with_math(value):
    segments = []
    if not value:
        return segments
    last = 0
    for match in MATH_PATTERN.finditer(value):
        start_idx, end_idx = match.span()
        if start_idx > last:
            segments.append(('text', value[last:start_idx], False))
        expr = match.group()
        display = False
        content = expr
        if expr.startswith('$$') and expr.endswith('$$'):
            content = expr[2:-2]
            display = True
        elif expr.startswith('$') and expr.endswith('$'):
            content = expr[1:-1]
            display = False
        elif expr.startswith(r'\[') and expr.endswith(r'\]'):
            content = expr[2:-2]
            display = True
        elif expr.startswith(r'\(') and expr.endswith(r'\)'):
            content = expr[2:-2]
            display = False
        content = content.strip()
        if content:
            segments.append(('math', content, display))
        last = end_idx
    if last < len(value):
        segments.append(('text', value[last:], False))
    return segments





def generate_math_assets(latex_expression, display=False, dpi=300):
    expression = latex_expression.strip()
    if expression.startswith('$$') and expression.endswith('$$') and len(expression) > 4:
        expression = expression[2:-2].strip()
    elif expression.startswith('$') and expression.endswith('$') and len(expression) > 2:
        expression = expression[1:-1].strip()
    if expression.startswith('\\(') and expression.endswith('\\)') and len(expression) > 4:
        expression = expression[2:-2].strip()
    if expression.startswith('\\[') and expression.endswith('\\]') and len(expression) > 4:
        expression = expression[2:-2].strip()
    if not expression:
        return None
    font_size = 16
    prop = FontProperties(size=font_size, family='STIXGeneral')
    math_string = f'$' + expression + '$'

    png_buffer = io.BytesIO()
    math_to_image(math_string, png_buffer, dpi=dpi, format='png', prop=prop, color='black')
    png_buffer.seek(0)
    with PILImage.open(png_buffer) as img:
        width_px, height_px = img.size
    png_buffer.seek(0)
    width_points = width_px * 72 / dpi
    height_points = height_px * 72 / dpi

    drawing = None
    if svg2rlg is not None:
        svg_buffer = io.BytesIO()
        try:
            math_to_image(math_string, svg_buffer, dpi=dpi, format='svg', prop=prop, color='black')
            svg_buffer.seek(0)
            drawing = svg2rlg(svg_buffer)
            if getattr(drawing, 'width', 0) and getattr(drawing, 'height', 0):
                scale_x = width_points / drawing.width
                scale_y = height_points / drawing.height
                scale = min(scale_x, scale_y)
                drawing.scale(scale, scale)
                drawing.width *= scale
                drawing.height *= scale
                drawing.hAlign = 'LEFT'
            else:
                drawing = None
        except Exception:
            drawing = None
        finally:
            svg_buffer.close()
    return png_buffer, width_points, height_points, drawing

def render_math_to_image(latex_expression, display=False, dpi=300):
    assets = generate_math_assets(latex_expression, display=display, dpi=dpi)
    if not assets:
        return None
    png_buffer, width_points, height_points, _ = assets
    png_buffer.seek(0)
    return png_buffer, width_points, height_points












def normalize_latex_spacing(text):
    if not text:
        return text
    text = str(text)

    def _collapse_display(match):
        inner = match.group(1).strip()
        inner = ' '.join(inner.split())
        return r'\[' + inner + r'\]'

    text = re.sub(r'\\\[\s*(.*?)\s*\\\]', _collapse_display, text, flags=re.S)

    def _collapse_inline(match):
        inner = match.group(1).strip()
        inner = ' '.join(inner.split())
        return r'\(' + inner + r'\)'

    text = re.sub(r'\\\(\s*(.*?)\s*\\\)', _collapse_inline, text, flags=re.S)

    text = re.sub(r'\\\s+([\\()\\[\\]])', lambda m: '\\' + m.group(1), text)
    text = re.sub(r'\\\s+([A-Za-z]+)', r'\\\1', text)
    return text

def append_text_with_math_to_story(story, text, style):
    if text is None:
        return
    text = strip_step_markers(text)
    lines = str(text).splitlines() or ['']
    for line in lines:
        line = normalize_latex_spacing(line)
        segments = split_text_with_math(line)
        if not segments:
            story.append(Spacer(1, 6))
            continue
        flow_items = []
        for kind, value, display in segments:
            if kind == 'text':
                clean_text = value.strip()
                if clean_text:
                    flow_items.append(Paragraph(escape(clean_text), style))
            elif kind == 'math':
                assets = generate_math_assets(value, display=display)
                if assets:
                    png_buffer, width_pt, height_pt, drawing = assets
                    if drawing is not None:
                        print('[debug] using vector drawing for PDF')
                        flow_items.append(renderPDF.GraphicsFlowable(drawing))
                    else:
                        print('[debug] falling back to PNG for PDF')
                        png_buffer.seek(0)
                        flow_items.append(RLImage(png_buffer, width=width_pt, height=height_pt))
                else:
                    print('[debug] no assets generated for expression')
        if flow_items:
            story.append(KeepTogether(flow_items))
            story.append(Spacer(1, 6))


def add_paragraph_with_math(document, text, style_name=None):
    if text is None:
        return
    text = strip_step_markers(text)
    lines = str(text).splitlines() or ['']
    for line in lines:
        line = normalize_latex_spacing(line)
        segments = split_text_with_math(line)
        if not segments:
            continue
        paragraph = document.add_paragraph()
        if style_name:
            paragraph.style = style_name
        for kind, value, display in segments:
            if kind == 'text':
                clean_text = value.strip()
                if clean_text:
                    paragraph.add_run(clean_text)
            elif kind == 'math':
                rendered = render_math_to_image(value, display=display)
                if rendered:
                    img_buffer, width_pt, height_pt = rendered
                    if display:
                        paragraph.alignment = 1
                    run = paragraph.add_run()
                    run.add_picture(img_buffer, width=Pt(width_pt * 0.9))
def strip_step_markers(text):
    if not text:
        return text
    filtered_lines = []
    for line in str(text).splitlines():
        stripped = line.strip()
        if stripped.startswith('主要ステップ') or stripped.startswith('最終解'):
            continue
        filtered_lines.append(line)
    return '\n'.join(filtered_lines).strip()


def parse_structured_sections(raw_text):
    sections = {}
    if not raw_text:
        return sections
    current_label = None
    for raw_line in str(raw_text).splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = re.match(r'^([^\s:：]{1,40})\s*[:：]\s*(.*)$', line)
        if match:
            current_label = match.group(1).strip()
            sections[current_label] = match.group(2).strip()
        elif current_label:
            sections[current_label] = (sections[current_label] + '\n' + line).strip()
    return sections


def parse_analysis_output(raw_text, problem_text):
    parsed = {
        'grade': None,
        'unit': None,
        'difficulty': None,
        'justification': None,
        'summary': None,
        'next_action_prompt': None,
    }
    normalized = ''
    if isinstance(raw_text, str):
        normalized = raw_text.strip()
    elif raw_text is None:
        normalized = ''
    else:
        normalized = str(raw_text).strip()

    payload = None
    if normalized:
        try:
            payload = json.loads(normalized)
        except json.JSONDecodeError:
            payload = None
    if isinstance(payload, dict):
        for key in parsed.keys():
            value = payload.get(key)
            if value:
                parsed[key] = str(value).strip()
        parsed['raw_text'] = str(payload.get('raw_text', normalized)).strip()
        problem_from_payload = payload.get('problem_text') or payload.get('original_problem')
        if problem_from_payload:
            parsed['problem_text'] = str(problem_from_payload).strip()
            parsed['original_problem'] = parsed['problem_text']
    else:
        sections = parse_structured_sections(normalized)
        label_map = {
            '学年': 'grade',
            '対象学年': 'grade',
            '単元': 'unit',
            'テーマ': 'unit',
            '分野': 'unit',
            '難易度': 'difficulty',
            'レベル': 'difficulty',
            '推定根拠': 'justification',
            '根拠': 'justification',
            '理由': 'justification',
            '要約': 'summary',
            '概要': 'summary',
            'まとめ': 'summary',
            '次のステップ': 'next_action_prompt',
            '次に学ぶ内容': 'next_action_prompt',
            '学習ステップ': 'next_action_prompt',
            '指導ポイント': 'next_action_prompt',
        }
        for label, value in sections.items():
            key = label_map.get(label)
            if key:
                parsed[key] = value
        if not parsed['justification'] and sections.get('推定根拠'):
            parsed['justification'] = sections['推定根拠']
        parsed['raw_text'] = normalized
    if 'problem_text' not in parsed or not parsed['problem_text']:
        parsed['problem_text'] = problem_text.strip() if problem_text else None
    if 'original_problem' not in parsed or not parsed['original_problem']:
        parsed['original_problem'] = parsed.get('problem_text')
    return parsed








def parse_generated_problems(raw_text):
    if not raw_text:
        return []
    text = str(raw_text).strip()
    if not text:
        return []

    problem_pattern = re.compile(
        r'(?:^|\n)\s*(?:【\s*)?問題\s*(\d+)\s*(?:】|[:：]|[.)．])',
        re.IGNORECASE,
    )
    matches = list(problem_pattern.finditer(text))
    if not matches:
        return []

    problems = []
    answer_pattern = re.compile(
        r'(?:【\s*解答\s*(\d+)\s*】|解答\s*(\d+)\s*(?:[:：]|[.)．]))',
        re.IGNORECASE,
    )
    explanation_pattern = re.compile(
        r'(?:【\s*解説\s*(\d+)\s*】|解説\s*(\d+)\s*(?:[:：]|[.)．]))',
        re.IGNORECASE,
    )

    for idx, match in enumerate(matches):
        number = int(match.group(1)) if match.group(1) else idx + 1
        start_idx = match.end()
        end_idx = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        block = text[start_idx:end_idx].strip()

        answer_body = None
        explanation_body = None
        problem_body = block

        answer_match = answer_pattern.search(block)
        if answer_match:
            problem_body = block[:answer_match.start()].strip()
            remainder = block[answer_match.end():].strip()
            explanation_match = explanation_pattern.search(remainder)
            if explanation_match:
                answer_body = remainder[:explanation_match.start()].strip() or None
                explanation_body = remainder[explanation_match.end():].strip() or None
            else:
                answer_body = remainder or None
        else:
            explanation_match = explanation_pattern.search(block)
            if explanation_match:
                problem_body = block[:explanation_match.start()].strip()
                explanation_body = block[explanation_match.end():].strip() or None

        title = None
        title_match = re.match(r'^(?:タイトル|題名)[:：]\s*(.*)$', problem_body)
        if title_match:
            title = title_match.group(1).strip()
            problem_body = problem_body[title_match.end():].strip()

        problem_body = normalize_latex_spacing(problem_body) if problem_body else None
        answer_body = normalize_latex_spacing(answer_body) if answer_body else None
        explanation_body = normalize_latex_spacing(explanation_body) if explanation_body else None

        problem_body = strip_step_markers(problem_body) if problem_body else None
        answer_body = strip_step_markers(answer_body) if answer_body else None
        explanation_body = strip_step_markers(explanation_body) if explanation_body else None

        problems.append({
            'number': number,
            'title': title,
            'problem': problem_body or None,
            'answer': answer_body,
            'explanation': explanation_body,
        })

    return problems

def build_problems_text(items):
    if not items:
        return ''
    lines = []
    for idx, item in enumerate(items, start=1):
        lines.append(f'【問題{idx}】')
        title = item.get('title')
        if title:
            lines.append(str(title).strip())
        problem_text = item.get('problem')
        if problem_text:
            lines.append(str(problem_text).strip())
        answer_text = item.get('answer')
        if answer_text:
            lines.append(f'【解答{idx}】')
            lines.append(str(answer_text).strip())
        explanation_text = item.get('explanation')
        if explanation_text:
            lines.append(f'【解説{idx}】')
            lines.append(str(explanation_text).strip())
        lines.append('')
    return '\n'.join(lines).strip()


def normalize_problems_text(raw_text):
    if not raw_text:
        return ''
    result = str(raw_text).strip()
    result = result.replace('�B��v�X�e�b�v', '�B\n��v�X�e�b�v')
    result = result.replace('�B\n�ŏI�� ��', '\n�ŏI�� ��')
    lines = [line for line in result.splitlines() if line.strip() != '�B']
    return strip_step_markers('\n'.join(lines))




@math_bp.route('/analyze', methods=['POST'])
@math_bp.route('/analyze-problem', methods=['POST'])
def analyze_problem():
    """例題を解析して単元と学年を推定"""
    try:
        data = request.get_json() or {}
        problem_text = (data.get('problem_text') or '').strip()

        if not problem_text:
            return jsonify({'error': '問題文が入力されていません'}), 400

        try:
            prompt_template = load_prompt_template()
        except Exception as e:
            print(f"プロンプトテンプレート読み込みエラー: {e}")
            return jsonify({'error': 'プロンプトテンプレートの読み込みに失敗しました'}), 500

        analysis_prompt = f"""
{prompt_template}

以下の例題を解析してください。解答は求めず、次の項目だけを順番を変えずに日本語で出力してください。

例題：
{problem_text}

出力形式（見出し名を変更しないこと）：
学年: [中1/中2/中3/数I/数A/数II/数B/数III/数C]
単元: [具体的な単元名]
難易度: [Level 1 (基礎) 〜 Level 5 (難関) などの表記]
推定根拠: [簡潔な説明]
要約: [解法の要点を1〜2文で]
次のステップ: [学習者への次の学習提案]
"""

        try:
            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "あなたは数学教育の専門家です。"},
                    {"role": "user", "content": analysis_prompt}
                ],
                max_tokens=500,
                temperature=0.3
            )
        except APIConnectionError:
            return jsonify({'error': 'OpenAI APIへの接続に失敗しました。ネットワーク環境とAPIキーを確認してください。'}), 503

        analysis_raw = (response.choices[0].message.content or '').strip()
        analysis_data = parse_analysis_output(analysis_raw, problem_text)
        analysis_data.setdefault('problem_text', problem_text)
        analysis_data.setdefault('original_problem', problem_text)

        return jsonify({
            'success': True,
            'analysis': analysis_data,
            'original_problem': problem_text,
            'raw_response': analysis_raw,
        })

    except Exception as e:
        print(f"例題解析エラー: {e}")
        traceback.print_exc()
        return jsonify({'error': f'解析中にエラーが発生しました: {str(e)}'}), 500



@math_bp.route('/generate', methods=['POST'])
@math_bp.route('/generate-problems', methods=['POST'])
def generate_problems():
    """類題を生成"""
    try:
        data = request.get_json() or {}
        analysis_data = data.get('analysis') or {}
        original_problem = (data.get('original_problem') or analysis_data.get('problem_text') or '').strip()

        if not original_problem:
            return jsonify({'error': '元の問題が指定されていません'}), 400

        difficulty = data.get('difficulty') or analysis_data.get('difficulty') or 'Level 3'
        count_raw = data.get('count') or analysis_data.get('count') or 3
        try:
            count = max(1, int(count_raw))
        except (TypeError, ValueError):
            count = 3

        solution_method = data.get('solution_method') or data.get('solution_hint') or ''
        analysis_summary = data.get('analysis_summary') or analysis_data.get('summary') or ''

        try:
            prompt_template = load_prompt_template()
        except Exception as e:
            print(f"プロンプトテンプレート読み込みエラー: {e}")
            return jsonify({'error': 'プロンプトテンプレートの読み込みに失敗しました'}), 500

        context_lines = []
        if analysis_data.get('grade'):
            context_lines.append(f"学年: {analysis_data['grade']}")
        if analysis_data.get('unit'):
            context_lines.append(f"単元: {analysis_data['unit']}")
        if analysis_data.get('difficulty'):
            context_lines.append(f"解析難易度: {analysis_data['difficulty']}")
        if analysis_summary:
            context_lines.append(f"要約: {analysis_summary}")
        if analysis_data.get('justification'):
            context_lines.append(f"推定根拠: {analysis_data['justification']}")
        if solution_method:
            context_lines.append(f"解法指定: {solution_method}")

        context_text = '\n'.join(context_lines) or '追加情報なし'

        generation_prompt = f"""
{prompt_template}

以下の例題と解析情報をもとに、新しい数学の類題を{count}問作成してください。

解析情報:
{context_text}

例題：
{original_problem}

出力形式（必ずこの形式を守ること）：
【問題1】
問題文
【解答1】
解答
【解説1】
解説
（指定した問題数になるまで番号を増やして繰り返す）
"""

        generation_prompt += """
追加ルール:
- 指定された作問数 {count} 問を必ず生成してください。
- ユーザーへの質問や確認は行わず、指定の形式のみで回答してください。
- 各問題には必ず問題文・解答・解説を含めてください。
"""

        if analysis_summary:
            generation_prompt += '\n各問題の解説には学習の要点を1文以上含めてください。'

        try:
            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "あなたは数学問題作成の専門家です。"},
                    {"role": "user", "content": generation_prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
        except APIConnectionError:
            return jsonify({'error': 'OpenAI APIへの接続に失敗しました。ネットワーク環境とAPIキーを確認してください。'}), 503

        raw_output = (response.choices[0].message.content or '').strip()
        problems = parse_generated_problems(raw_output)
        problems_text = build_problems_text(problems) if problems else raw_output

        metadata = {
            'grade': analysis_data.get('grade'),
            'unit': analysis_data.get('unit'),
            'difficulty': difficulty,
        }

        notes_sections = []
        if analysis_data.get('justification'):
            notes_sections.append(str(analysis_data['justification']).strip())
        if solution_method:
            notes_sections.append(f"解法指定: {solution_method}")
        if analysis_summary:
            notes_sections.append(f"解析要約: {analysis_summary}")
        if notes_sections:
            metadata['notes'] = '\n'.join(notes_sections)

        return jsonify({
            'success': True,
            'problems': problems,
            'problems_text': problems_text,
            'metadata': metadata,
            'raw_response': raw_output,
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'類題生成中にエラーが発生しました: {str(e)}'}), 500



@math_bp.route('/ocr-image', methods=['POST'])
def ocr_image():
    """画像からテキストを抽出（OCR）"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': '画像ファイルが選択されていません'}), 400

        image_file = request.files['image']

        image_data = image_file.read()
        base64_image = base64.b64encode(image_data).decode('utf-8')

        try:
            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "この画像に含まれる数学問題のテキストを正確に読み取って、テキスト形式で出力してください。数式は適切な記法で表現してください。"
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=1000
            )
        except APIConnectionError:
            return jsonify({'error': 'OpenAI APIへの接続に失敗しました。ネットワーク環境とAPIキーを確認してください。'}), 503

        extracted_text = response.choices[0].message.content

        return jsonify({
            'success': True,
            'extracted_text': extracted_text
        })

    except Exception as e:
        return jsonify({'error': f'OCR処理中にエラーが発生しました: {str(e)}'}), 500


@math_bp.route('/download/pdf', methods=['POST'])
@math_bp.route('/export-pdf', methods=['POST'])
@math_bp.route('/download/pdf', methods=['POST'])
@math_bp.route('/export-pdf', methods=['POST'])
def export_pdf():
    """生成された問題をPDF形式でエクスポート"""
    try:
        data = request.get_json() or {}
        metadata = data.get('metadata') or {}
        problems = data.get('problems') or []
        problems_text = strip_step_markers(data.get('problems_text') or '')
        problems_text = normalize_latex_spacing(problems_text)

        if not problems:
            parsed = parse_generated_problems(problems_text)
            if parsed:
                problems = parsed

        if problems and not problems_text:
            problems_text = build_problems_text(problems)

        problems_text = strip_step_markers(normalize_problems_text(problems_text))
        problems_text = normalize_latex_spacing(problems_text)

        if not problems and not problems_text:
            return jsonify({'error': '出力する問題がありません'}), 400

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=DEFAULT_FONT_NAME,
            fontSize=16,
            spaceAfter=24,
        )

        section_title_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading2'],
            fontName=DEFAULT_FONT_NAME,
            fontSize=14,
            spaceAfter=12,
        )

        problem_heading_style = ParagraphStyle(
            'ProblemHeading',
            parent=styles['Heading3'],
            fontName=DEFAULT_FONT_NAME,
            fontSize=13,
            spaceAfter=6,
        )

        label_style = ParagraphStyle(
            'LabelStyle',
            parent=styles['Normal'],
            fontName=DEFAULT_FONT_NAME,
            fontSize=12,
            spaceAfter=4,
        )

        content_style = ParagraphStyle(
            'CustomContent',
            parent=styles['Normal'],
            fontName=DEFAULT_FONT_NAME,
            fontSize=12,
            spaceAfter=6,
        )

        story = []
        title_text = metadata.get('unit') or '数学問題集'
        story.append(Paragraph(title_text, title_style))
        story.append(Spacer(1, 12))

        problems_list = problems or parse_generated_problems(problems_text)
        if not problems_list:
            for line in problems_text.split('\n'):
                append_text_with_math_to_story(story, line, content_style)
        else:
            story.append(Paragraph('問題一覧', section_title_style))
            story.append(Spacer(1, 6))
            for idx, item in enumerate(problems_list, start=1):
                story.append(Paragraph(f'問題{idx}', problem_heading_style))
                append_text_with_math_to_story(story, item.get('problem'), content_style)
                story.append(Spacer(1, 12))

            story.append(PageBreak())
            story.append(Paragraph('解答・解説', section_title_style))
            story.append(Spacer(1, 6))
            for idx, item in enumerate(problems_list, start=1):
                story.append(Paragraph(f'問題{idx}', problem_heading_style))
                answer = item.get('answer')
                explanation = item.get('explanation')
                if answer:
                    story.append(Paragraph('解答', label_style))
                    append_text_with_math_to_story(story, answer, content_style)
                if explanation:
                    story.append(Paragraph('解説', label_style))
                    append_text_with_math_to_story(story, explanation, content_style)
                story.append(Spacer(1, 12))

        doc.build(story)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name='math_problems.pdf',
            mimetype='application/pdf'
        )

    except APIConnectionError:
        traceback.print_exc()
        return jsonify({'error': 'PDF出力中にエラーが発生しました: 接続エラー'}), 500


@math_bp.route('/download/word', methods=['POST'])
@math_bp.route('/export-word', methods=['POST'])
def export_word():
    """生成された問題をWord形式でエクスポート"""
    try:
        data = request.get_json() or {}
        metadata = data.get('metadata') or {}
        problems = data.get('problems') or []
        problems_text = strip_step_markers(data.get('problems_text') or '')
        problems_text = normalize_latex_spacing(problems_text)

        if not problems:
            parsed = parse_generated_problems(problems_text)
            if parsed:
                problems = parsed

        if problems and not problems_text:
            problems_text = build_problems_text(problems)

        if not problems and not problems_text:
            return jsonify({'error': '出力する問題がありません'}), 400

        doc = Document()
        title_text = metadata.get('unit') or '数学問題集'
        doc.add_heading(title_text, 0)

        problems_list = problems or parse_generated_problems(problems_text)
        if not problems_list:
            for line in strip_step_markers(problems_text or '').split('\n'):
                add_paragraph_with_math(doc, line)
        else:
            doc.add_heading('問題一覧', level=1)
            for idx, item in enumerate(problems_list, start=1):
                doc.add_heading(f'問題{idx}', level=2)
                add_paragraph_with_math(doc, item.get('problem'))
                doc.add_paragraph('')

            doc.add_page_break()
            doc.add_heading('解答・解説', level=1)
            for idx, item in enumerate(problems_list, start=1):
                doc.add_heading(f'問題{idx}', level=2)
                answer = item.get('answer')
                explanation = item.get('explanation')
                if answer:
                    doc.add_heading('解答', level=3)
                    add_paragraph_with_math(doc, answer)
                if explanation:
                    doc.add_heading('解説', level=3)
                    add_paragraph_with_math(doc, explanation)
                doc.add_paragraph('')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name='math_problems.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'Word出力中にエラーが発生しました: {str(e)}'}), 500














